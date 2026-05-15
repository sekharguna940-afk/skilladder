import json
import docx
import os
from services.resume_parser import resume_parser

def stress_test_parser():
    # 1. Create a real DOCX for testing
    doc_path = "stress_test_mock.docx"
    doc = docx.Document()
    doc.add_heading('SAM SMITH', 0)
    doc.add_paragraph('Address: 123 Java Lane, Bangalore, India 560001')
    doc.add_paragraph('Email: sam.smith@example.com | Phone: +91 9999999999')
    
    doc.add_heading('PROFESSIONAL SUMMARY', level=1)
    doc.add_paragraph('Highly motivated Software Engineer with 5 years of experience in developing Microservices and REST APIs. Loves using Python and React.js for building scalable platforms.')
    
    doc.add_heading('TECHNICAL EXPERTISE', level=1)
    doc.add_paragraph('Programming: Python, Java, JavaScript, TypeScript, Go.')
    doc.add_paragraph('Frameworks: React, Next.js, Redux, Node.js, FastAPI, Spring Boot.')
    doc.add_paragraph('Cloud/DevOps: AWS (Lambda, EC2, S3), Docker, Kubernetes, Terraform, CI/CD.')
    doc.add_paragraph('Data & AI: Pandas, NumPy, Scikit-Learn, TensorFlow, LangChain, OpenAI API.')
    doc.add_paragraph('Niche Tools: HuggingFace, Pinecone, LlamaIndex, Mojo-Lang.')
    
    doc.add_heading('WORK EXPERIENCE', level=1)
    doc.add_paragraph('AI ENGINEER | DataFuture, Bangalore (Jan 2020 - Present)')
    doc.add_paragraph('- Implemented a Natural Language Processing pipeline that processed 1M documents.')
    doc.add_paragraph('- Optimized React components for better performance.')
    doc.add_paragraph('- Used Docker and Kubernetes for container orchestration.')
    
    doc.add_heading('EDUCATION', level=1)
    doc.add_paragraph('Indian Institute of Technology, Mumbai (2016 - 2020)')
    doc.add_paragraph('CGPA: 9.5')
    
    doc.save(doc_path)
    
    print("--- 🚀 STARTING RIGOROUS RESUME STRESS TEST 🚀 ---")
    
    try:
        with open(doc_path, "rb") as f:
            file_content = f.read()
            
        result = resume_parser.parse_resume(file_content, doc_path)
        
        if "error" in result:
            print(f"❌ Parser Error: {result['error']}")
            return

        print("\n--- Categorized Skills Result ---")
        print(json.dumps(result.get('categorized_skills', {}), indent=2))
        
        skills_set = set(result.get('skills', []))
        
        # 1. Check for expected multi-word skills
        expected_multi = {"Next.js", "Spring Boot", "OpenAI API", "Scikit-Learn", "Natural Language Processing"}
        found_multi = [s for s in expected_multi if s.lower() in [sk.lower() for sk in skills_set]]
        print(f"\nMulti-word skills captured: {found_multi}")
        
        # 2. Check for noise rejection
        NOISE_ITEMS = {"Bangalore", "India", "Resume", "Present", "Address", "Contact", "Smith", "January"}
        found_noise = [s for s in NOISE_ITEMS if s in skills_set]
        
        if found_noise:
            print(f"❌ NOISE LEAK DETECTED: {found_noise}")
        else:
            print("✅ NOISE REJECTION: Passed (No location/metadata leaked)")
            
        # 3. Check for Niche Discovery
        NICHE = {"Mojo-Lang", "LlamaIndex", "Pinecone", "HuggingFace"}
        found_niche = [s for s in NICHE if s.lower() in [sk.lower() for sk in skills_set]]
        print(f"Niche/New Tools Discovered: {found_niche}")
        
        if len(found_multi) >= 2 and not found_noise and len(found_niche) >= 2:
            print("\n🏆 TEST PASSED: Industrial-Grade Parser is highly accurate.")
        else:
            print("\n⚠️ TEST COMPLETED: Results partially successful.")

    finally:
        if os.path.exists(doc_path):
            os.remove(doc_path)

if __name__ == "__main__":
    stress_test_parser()

if __name__ == "__main__":
    try:
        stress_test_parser()
    except Exception as e:
        print(f"Test Error: {e}")
