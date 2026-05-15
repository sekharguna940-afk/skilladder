import json
import docx
import os
import sys

# Add the current directory to path so we can import services
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), 'services')))
from services.resume_parser import resume_parser

def test_industrial_skills():
    # 1. Create a dummy resume with the user's provided skills
    doc_path = "industrial_test.docx"
    doc = docx.Document()
    doc.add_heading('INDUSTRIAL SPECIALIST', 0)
    
    doc.add_heading('KEY SKILLS & COMPETENCIES', level=1)
    skills_text = """
    Equipment Commissioning
    Site Maintenance & Troubleshooting
    Production Control & Operation Management
    EPLAN & AUTOCAD Drawing
    Material Management & Procurement
    Client Relationship Management
    PLC Programming
    Production Control
    Manpower Management
    Communicator
    Motivational Leader
    Strategic Thinker
    Team Player
    Innovative
    """
    doc.add_paragraph(skills_text)
    
    doc.add_heading('AUTOMATION & IT SKILLS', level=1)
    automation_text = """
    Design Software: AutoCAD & EPLAN
    Installation: PLC Projects (Wehrhann AAC Plant & ASRS Warehouse & Track and Trace System)
    PLC Skills: Siemens PLC (S7 300/1200), OMRON PLC, AB & ABB PLC
    Programming Language: Ladder Diagram & Functional Block Diagram
    VFD: SEW & Siemens (Parameter Setting & Keypad Programming)
    HMI & SCADA: Operations & Design using WinCC Software
    Software: Simatic Manager (Siemens), Ladder Builder for KV (Keyence), CX Programmer (OMRON)
    Communication: PROFIBUS PROFINET, ASI Communication & Addressing
    Industrial Equipment: Encoder, IR Device, Profibus Coupler, Contactor, Relays, Various Circuit Breaker, Bus Bar, VFD, Single & Three Phase Induction Motor & Various Types of Sensors
    Data Center: Installation of Server, Storage System, Network Switches, Firewall Configuration & RedHat, CentOS, Oracle SQL & MSSQL and Web Application Installation on Tomcat Web Server.
    """
    doc.add_paragraph(automation_text)
    
    doc.add_heading('ADVANCED AUTOMATION & INDUSTRY 4.0', level=1)
    advanced_text = """
    Advanced Systems: PID Control Tuning, Servo Motor & Motion Control, Root Cause Analysis (RCA)
    Industry 4.0: IIoT Integration, Digital Twin Concepts, OEE Monitoring, MES
    Network: Ethernet/IP Configuration, Industrial Cybersecurity Basics
    """
    doc.add_paragraph(advanced_text)
    
    doc.add_heading('PROJECT & PROFESSIONAL SKILLS', level=1)
    mgmt_text = """
    Methodologies: Lean Manufacturing, Kaizen, Six Sigma Basics (Green Belt)
    Project Management: PMP Certified, MS Project, Budget Planning, Risk Assessment
    Professional: Stakeholder Management, Change Management, Cross-Functional Coordination
    """
    doc.add_paragraph(mgmt_text)
    
    doc.add_heading('EDUCATION', level=1)
    edu_text = """
    Bachelor of Engineering (Mechanical)
    University of Engineering - 2018
    Academic Score: 8.5/10 CGPA (85% percentage)
    """
    doc.add_paragraph(edu_text)
    
    # Add noise samples provided by user
    noise_text = """
    rwand aproject
    (Rwanda)- On-site Project
    12 Helpers
    9 Technician
    AAC Factory Commissioning
    ABB PLC
    ASRS Project
    ASRS Warehouse
    ATKO
    ATKO Scales Pvt. Ltd.
    Bus Bar
    """
    doc.add_heading('PROJECT CONTEXT', level=1)
    doc.add_paragraph(noise_text)
    
    doc.save(doc_path)
    
    print("--- 🛠️ VERIFYING MASSIVE TAXONOMY & NUMERICAL ACCURACY 🛠️ ---")
    
    try:
        with open(doc_path, "rb") as f:
            file_content = f.read()
            
        result = resume_parser.parse_resume(file_content, doc_path)
        
        extracted_skills = set(result.get('skills', []))
        
        # New key skills to verify
        EXPANDED_SKILLS = [
            "PID Control Tuning", "IIoT", "Digital Twin", "OEE", "Lean Manufacturing",
            "Kaizen", "Six Sigma", "PMP", "MS Project", "Stakeholder Management",
            "Servo Motor", "RCA", "MES"
        ]
        
        found = []
        missing = []
        
        for s in EXPANDED_SKILLS:
            if any(s.lower() in es.lower() for es in extracted_skills):
                found.append(s)
            else:
                missing.append(s)
                
        # Verify Numerical Extraction
        detected_cgpa = result.get('cgpa')
        detected_perc = result.get('percentage')
        academic_score = result.get('academic_score')
        score_type = result.get('score_type')
        
        print(f"\n✅ Skills Found: {len(found)}/{len(EXPANDED_SKILLS)}")
        print(f"Detected: {', '.join(found[:10])}...")
        
        if missing:
            print(f"❌ Missing: {', '.join(missing)}")
            
        print("\n--- Numerical Extraction Check ---")
        if detected_cgpa == "8.5":
            print("✅ CGPA Extraction (8.5/10): Passed")
        else:
            print(f"❌ CGPA Extraction: Failed (Got {detected_cgpa})")
            
        if detected_perc == "85":
            print("✅ Percentage Extraction (85%): Passed")
        else:
            print(f"❌ Percentage Extraction: Failed (Got {detected_perc})")

        print("\n--- Unified Academic Score Check ---")
        if academic_score == "8.5" and score_type == "CGPA":
            print("✅ Unified Score (CGPA): Passed")
        elif academic_score == "85%" and score_type == "Percentage":
            print("✅ Unified Score (Percentage): Passed")
        else:
            print(f"❌ Unified Score: Failed (Got {academic_score} as {score_type})")

        # Check Categorization
        print("\n--- Categorization Check ---")
        cats = result.get('categorized_skills', {})
        
        target_cats = ["Project and Business Management", "Industry 4.0 and Smart Manufacturing", "Certifications"]
        for cat in target_cats:
            if cats.get(cat):
                print(f"✅ {cat}: {cats[cat][:3]}")
            else:
                print(f"❌ {cat}: Missing (Skills found in this cat: {[s for s in result.get('skills', []) if s in resume_parser.taxonomy.get(cat, [])]})")

        if len(found) > 10 and detected_cgpa == "8.5":
            print("\n🏆 ULTIMATE VERIFICATION SUCCESS: Massive taxonomy and accuracy boosted!")
        else:
            print("\n⚠️ VERIFICATION INCOMPLETE.")

    finally:
        if os.path.exists(doc_path):
            os.remove(doc_path)

if __name__ == "__main__":
    test_industrial_skills()
